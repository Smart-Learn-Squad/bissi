"""Word document operations for BISSI.

Provides DocxAgent and StyledDocument classes for reading and creating Microsoft Word documents.

Classes:
- DocxAgent: Read, manipulate, and save .docx files
- StyledDocument: Create styled documents with titles, headings, tables, lists, images

Functions:
- read_document: Read all text from a Word document
- read_with_structure: Read document with paragraphs and tables
- create_document: Create a new empty Word document
- write_document: Write text content to a Word document
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from typing import Optional, Union, List


# Predefined color palettes for consistent styling
class DocColors:
    """Predefined color palette for BISSI documents."""
    PRIMARY = RGBColor(46, 134, 171)      # #2E86AB - Deep blue
    SECONDARY = RGBColor(33, 33, 33)       # #212121 - Dark gray
    ACCENT = RGBColor(231, 76, 60)       # #E74C3C - Coral red
    SUCCESS = RGBColor(39, 174, 96)       # #27AE60 - Green
    WARNING = RGBColor(241, 196, 15)      # #F1C40F - Yellow
    INFO = RGBColor(52, 152, 219)       # #3498DB - Light blue
    LIGHT_GRAY = RGBColor(149, 165, 166)  # #95A5A6
    WHITE = RGBColor(255, 255, 255)
    BLACK = RGBColor(0, 0, 0)


# Table style presets
class TableStyles:
    """Predefined table styles."""
    GRID = 'Table Grid'
    LIGHT_ACCENT = 'Light List Accent 1'
    DARK_GRID = 'Dark Grid'


class DocxAgent:
    """A simple agent to read and manipulate docx documents."""

    def __init__(self, file_path: str):
        """Initialize the agent with the path to the docx file.

        Args:
            file_path: Path to the .docx file
        """
        self.file_path = file_path
        self.document = Document(file_path)

    def read_paragraphs(self) -> list[str]:
        """Read and return the text of all paragraphs in the document.

        Returns:
            List of paragraph texts
        """
        return [para.text for para in self.document.paragraphs]

    def read_tables(self) -> list[list[list[str]]]:
        """Read and return the data from all tables in the document.

        Returns:
            List of tables, each table is a list of rows, each row is a list of cell texts
        """
        tables_data = []
        for table in self.document.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables_data.append(table_data)
        return tables_data

    def save(self, new_file_path: str = None) -> None:
        """Save the document.

        Args:
            new_file_path: Optional new path to save to; if None, overwrite original
        """
        if new_file_path:
            self.document.save(new_file_path)
        else:
            self.document.save(self.file_path)


# Module-level convenience functions
def read_document(file_path: str) -> str:
    """Read all text from a Word document."""
    agent = DocxAgent(file_path)
    return "\n".join(agent.read_paragraphs())


def read_with_structure(file_path: str) -> dict:
    """Read document with structure (paragraphs and tables)."""
    agent = DocxAgent(file_path)
    return {
        "paragraphs": agent.read_paragraphs(),
        "tables": agent.read_tables()
    }


def create_document() -> DocxAgent:
    """Create a new empty Word document."""
    from docx import Document
    from pathlib import Path
    import tempfile

    # Create temp file
    temp_path = Path(tempfile.gettempdir()) / "new_document.docx"
    doc = Document()
    doc.save(str(temp_path))
    return DocxAgent(str(temp_path))


def write_document(file_path: str, content: str, append: bool = False) -> bool:
    """Write text content to a Word document.

    Args:
        file_path: Path to the .docx file
        content: Text to write (can contain newlines for paragraphs)
        append: If True, add to end of document. If False, create new/overwrite.
    """
    from docx import Document
    import os

    if append and os.path.exists(file_path):
        doc = Document(file_path)
    else:
        doc = Document()

    # Split content by newlines and add as paragraphs
    for line in content.split('\n'):
        if line.strip():
            doc.add_paragraph(line)

    doc.save(file_path)
    return True


class StyledDocument:
    """Create polished, styled Word documents with headings, tables, lists, and images.

    Provides a fluent API for professional document generation with:
    - Title and heading levels (1-4)
    - Styled paragraphs (bold, italic, colors, font sizes)
    - Tables with header rows and styling
    - Bulleted/numbered lists
    - Images and inline graphics
    - Page setup (margins, orientation)

    Example:
        doc = StyledDocument("rapport.docx")
        doc.add_title("Annual Report 2024")
        doc.add_heading("Executive Summary")
        doc.add_paragraph("This report outlines...", bold=True)
        doc.add_table(data, header_row=True, style=TableStyles.GRID)
        doc.add_image("chart.png", width=Cm(15))
        doc.save()
    """

    def __init__(
        self,
        file_path: Optional[str] = None,
        margin_top: float = 2.0,
        margin_bottom: float = 2.0,
        margin_left: float = 2.5,
        margin_right: float = 2.5,
    ):
        """Initialize a styled document.

        Args:
            file_path: Path to save the document (optional, can set later)
            margin_top: Top margin in cm (default: 2.0)
            margin_bottom: Bottom margin in cm (default: 2.0)
            margin_left: Left margin in cm (default: 2.5)
            margin_right: Right margin in cm (default: 2.5)
        """
        self.file_path = file_path
        self.document = Document() if not file_path or not Path(file_path).exists() else Document(file_path)

        # Configure page margins
        section = self.document.sections[0]
        section.top_margin = Cm(margin_top)
        section.bottom_margin = Cm(margin_bottom)
        section.left_margin = Cm(margin_left)
        section.right_margin = Cm(margin_right)

        self._default_font = "Calibri"
        self._default_font_size = 11

    # ─────────────────────────────────────────────────────────
    # Title & Headings
    # ─────────────────────────────────────────────────────────

    def add_title(self, text: str, level: int = 0, color: RGBColor = None) -> 'StyledDocument':
        """Add a document title (centered, large).

        Args:
            text: Title text
            level: 0 = Title (36pt), 1 = Heading 1 (24pt), etc.
            color: RGBColor from DocColors or custom

        Returns:
            self for fluent chaining
        """
        sizes = {0: 36, 1: 28, 2: 24, 3: 20, 4: 16}
        font_size = sizes.get(level, 36)
        font_color = color or DocColors.PRIMARY

        if level == 0:
            # Title - centered, large
            paragraph = self.document.add_heading(text, level=0)
            for run in paragraph.runs:
                run.font.size = Pt(font_size)
                run.font.bold = True
                run.font.color.rgb = font_color
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Heading levels
            heading = self.document.add_heading(text, level=level)
            for run in heading.runs:
                run.font.size = Pt(font_size)
                run.font.bold = True
                run.font.color.rgb = font_color

        return self

    def add_heading(self, text: str, level: int = 1, color: RGBColor = None) -> 'StyledDocument':
        """Add a heading with specified level.

        Args:
            text: Heading text
            level: 1-4 (1=largest)
            color: RGBColor

        Returns:
            self for fluent chaining
        """
        return self.add_title(text, level=level, color=color)

    # ─────────────────────────────────────────────────────────
    # Paragraphs
    # ─────────────────────────────────────────────────────────

    def add_paragraph(
        self,
        text: str,
        bold: bool = False,
        italic: bool = False,
        font_size: int = None,
        color: RGBColor = None,
        alignment: str = 'left',
        space_before: float = 0,
        space_after: float = 6,
    ) -> 'StyledDocument':
        """Add a styled paragraph.

        Args:
            text: Paragraph text
            bold: Bold text
            italic: Italic text
            font_size: Font size in points (default: document default)
            color: RGBColor
            alignment: 'left', 'center', 'right', 'justify'
            space_before: Space before paragraph in points
            space_after: Space after paragraph in points

        Returns:
            self for fluent chaining
        """
        paragraph = self.document.add_paragraph()

        # Configure spacing
        paragraph.space_before = Pt(space_before) if space_before else None
        paragraph.space_after = Pt(space_after) if space_after else None

        # Configure alignment
        align_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        paragraph.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

        # Build run with styles
        run = paragraph.add_run(text)
        run.font.name = self._default_font
        run.font.size = Pt(font_size or self._default_font_size)

        if bold:
            run.font.bold = True
        if italic:
            run.font.italic = True
        if color:
            run.font.color.rgb = color

        return self

    def add_spacer(self, height: float = 12) -> 'StyledDocument':
        """Add vertical space between elements.

        Args:
            height: Space height in points

        Returns:
            self for fluent chaining
        """
        self.document.add_paragraph().space_before = Pt(height)
        return self

    # ─────────────────────────────────────────────────────────
    # Lists
    # ─────────────────────────────────────────────────────────

    def add_bullet_list(self, items: List[str], bold_first: bool = False) -> 'StyledDocument':
        """Add a bulleted list.

        Args:
            items: List items to add
            bold_first: Make first item bold

        Returns:
            self for fluent chaining
        """
        for i, item in enumerate(items):
            p = self.document.add_paragraph(item, style='List Bullet')
            if i == 0 and bold_first:
                for run in p.runs:
                    run.font.bold = True
        return self

    def add_numbered_list(self, items: List[str], start: int = 1) -> 'StyledDocument':
        """Add a numbered list.

        Args:
            items: List items to add
            start: Starting number

        Returns:
            self for fluent chaining
        """
        for i, item in enumerate(items, start):
            p = self.document.add_paragraph(f"{i}. {item}", style='List Number')
        return self

    # ─────────────────────────────────────────────────────────
    # Tables
    # ─────────────────────────────────────────────────────────

    def add_table(
        self,
        data: List[List[str]],
        header_row: bool = True,
        style: str = None,
        column_widths: List[float] = None,
        bold_headers: bool = True,
        header_color: RGBColor = DocColors.PRIMARY,
    ) -> 'StyledDocument':
        """Add a styled table.

        Args:
            data: 2D list of cell values [[row1], [row2], ...]
            header_row: First row is header
            style: Table style name from TableStyles
            column_widths: Column widths in cm (optional)
            bold_headers: Make header row bold
            header_color: Header text color

        Returns:
            self for fluent chaining
        """
        if not data:
            return self

        table = self.document.add_table(rows=len(data), cols=len(data[0]))
        table.style = style

        for i, row_data in enumerate(data):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(cell_text)

                # Style header row
                if i == 0 and header_row:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = header_color
                            run.font.size = Pt(self._default_font_size - 1)
                    # Header background via XML
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'E8F4F8')  # Light blue background
                    cell._tc.get_or_add_tcPr().append(shading_elm)

        # Set column widths
        if column_widths and len(column_widths) == len(data[0]):
            for i, width in enumerate(column_widths):
                for row in table.rows:
                    row.cells[i].width = Cm(width)

        return self

    def add_data_table(
        self,
        rows: List[dict],
        columns: List[str] = None,
        header_row: bool = True,
    ) -> 'StyledDocument':
        """Add a table from list of dictionaries.

        Args:
            rows: [{col1: val, col2: val}, ...]
            columns: Column order (uses dict keys if None)
            header_row: Include header row

        Returns:
            self for fluent chaining
        """
        if not rows:
            return self

        # Determine columns
        cols = columns or list(rows[0].keys())
        data = []

        if header_row:
            data.append(cols)

        for row in rows:
            data.append([str(row.get(col, '')) for col in cols])

        return self.add_table(data, header_row=header_row)

    # ─────────────────────────────────────────────────────────
    # Images
    # ─────────────────────────────────────────────────────────

    def add_image(
        self,
        image_path: str,
        width: float = None,
        height: float = None,
        align: str = 'center',
    ) -> 'StyledDocument':
        """Add an image to the document.

        Args:
            image_path: Path to image file
            width: Width in cm (maintains aspect ratio if height not set)
            height: Height in cm
            align: 'left', 'center', 'right'

        Returns:
            self for fluent chaining
        """
        paragraph = self.document.add_paragraph()

        # Add image
        run = paragraph.add_run()
        run.add_picture(image_path, width=Cm(width) if width else None, height=Cm(height) if height else None)

        # Alignment
        align_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
        }
        paragraph.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.CENTER)

        return self

    def add_image_paragraph(self, image_path: str, caption: str = None) -> 'StyledDocument':
        """Add image with optional caption below.

        Args:
            image_path: Path to image
            caption: Caption text below image

        Returns:
            self for fluent chaining
        """
        self.add_image(image_path)

        if caption:
            self.add_paragraph(caption, italic=True, font_size=10, alignment='center')

        return self

    # ─────────────────────────────────────────────────────────
    # Page Layout
    # ─────────────────────────────────────────────────────────

    def set_page_orientation(self, orientation: str = 'portrait') -> 'StyledDocument':
        """Set page orientation.

        Args:
            orientation: 'portrait' or 'landscape'

        Returns:
            self for fluent chaining
        """
        section = self.document.sections[0]
        if orientation == 'landscape':
            section.orientation = 1  # WD_ORIENT.LANDSCAPE
        else:
            section.orientation = 0  # WD_ORIENT.PORTRAIT
        return self

    def set_page_size(self, width_cm: float = 21.0, height_cm: float = 29.7) -> 'StyledDocument':
        """Set custom page size (A4 = 21×29.7).

        Args:
            width_cm: Page width in cm
            height_cm: Page height in cm

        Returns:
            self for fluent chaining
        """
        section = self.document.sections[0]
        section.page_width = Cm(width_cm)
        section.page_height = Cm(height_cm)
        return self

    # ─────────────────────────────────────────────────────────
    # Save
    # ─────────────────────────────────────────────────────────

    def save(self, file_path: str = None) -> None:
        """Save the document.

        Args:
            file_path: Path to save (uses init path if None)
        """
        path = file_path or self.file_path
        if not path:
            raise ValueError("No file_path specified")
        self.document.save(path)

    def get_document(self) -> Document:
        """Get underlying python-docx Document object.

        Returns:
            Document
        """
        return self.document